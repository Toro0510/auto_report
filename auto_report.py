# -*- coding: utf-8 -*-
"""
Created on Wed Aug 14 16:24:01 2019

@author: jizeyuan
"""

from pyhive import hive
from sqlalchemy import create_engine
import pandas as pd
from docx import Document


conn_hive = hive.Connection(host='47.96.69.25', port=21050, username='admin', auth='NOSASL')
conn_sql = create_engine(
    'mysql+mysqldb://)



def loop_trade_no(sql_code, TRADE_NO):
    return sql_code.format(RISK_TRADE_NO='\'' + str(TRADE_NO) + '\'')



hive_base_info = pd.read_sql(loop_trade_no(
    '''
    select 
    a.trade_no,
    a.user_id,
    b.risk_user_id,
    a.dt_created,
    e.approve_status,
    f.name,
    f.regphone,
    f.sendphone,
    f.`identi_card`,
    floor(datediff(from_unixtime(unix_timestamp(),'yyyy-MM-dd'),
    concat_ws('-',substr(f.`identi_card`,7,4),substr(f.`identi_card`,11,2),substr(f.`identi_card`,13,2)))/365) as age,
    IF(substr(f.`identi_card`,17,1) IN ('1','3','5','7','9'),'M','F') AS gender,
    if(d.is_new_user_submmit=1,'新户','旧户') as is_new_user,
    (CASE WHEN a.channel_id=9 THEN '花呗' 
          WHEN a.channel_id=10 THEN '芝麻' 
          WHEN a.channel_id=19 THEN 'ios' 
          WHEN a.channel_id=22 THEN '安卓' 
          WHEN a.channel_id IN (31,33,34,36,37,39,46,50) THEN '小程序' 
          ELSE '其他' END) AS channel_name,
    if(a.order_type=1,'长租','短租') as order_type,
    c.sku_name,
    c.phone_type,
    a.freeze_cash,
    a.cash,
    round(a.freeze_cash/a.cash,4) as freeze_rate,
    a.pre_installments_num,
    a.pre_installments_total,
    g.address as send_address
    from dw_2.order_all a
    left join dw_2.user_id_relate_risk b on a.user_id=b.user_id --拼接risk_user_id
    left join dw_2.product_sku_main c on a.id_sku=c.sku_id --SKU_name
    left join dw_2.c_order_lab d on a.trade_no=d.trade_no --新旧户
    left join dw_2.risk_after_sale_approve_result e on a.trade_no=e.trade_no --审核状态
    left join dw.user_detail f on a.user_id=f.id_user --客户基本信息
    left join stage.airentdbsetting_airent_tbl_trade_delivery_address g on a.trade_no=g.trade_no --收货地址
    where a.trade_no={RISK_TRADE_NO}
    '''
    ,
    , conn_hive)


sql_repay_info = pd.read_sql('''
                           SELECT installment,debit_date,finished_time,overdue_days,is_finished,is_delete FROM `dw_risk`.`risk_payplan`
                           WHERE main_trade_no={RISK_TRADE_NO}
                           '''
                             , conn_sql)


collect_note = pd.read_sql('''
                           SELECT call_employee_name,call_time,call_result FROM dw_risk.risk_calllog
                           WHERE trade_no='20190505100230101434'
                           '''
                           , conn_sql)


approve_log = pd.read_sql('''
                           SELECT approver_name,remark,create_time FROM ods.`riskcontrol_e_approve`
                           WHERE TRADE_NO='20190505100230101434'
                           '''
                          , conn_sql)


application_log = pd.read_sql('''
                            select 
                            bb.dt_created,
                            --bb.trade_no,bb.risk_user_id,bb.user_id,
                            (CASE WHEN bb.channel_id=9 THEN '花呗' 
                                  WHEN bb.channel_id=10 THEN '芝麻' 
                                  WHEN bb.channel_id=19 THEN 'ios' 
                                  WHEN bb.channel_id=22 THEN '安卓' 
                                  WHEN bb.channel_id IN (31,33,34,36,37,39,46,50) THEN '小程序' 
                                  ELSE '其他' END) AS channel_name,
                            if(bb.order_type=1,'长租','短租') as order_type,
                            --cc.sku_name,
                            cc.phone_type,
                            bb.freeze_cash,
                            bb.cash,
                            (CASE WHEN ff.STATUS=0 THEN '初始状态'
                                  WHEN ff.STATUS=1 THEN '待支付（订单提交）'
                                  WHEN ff.STATUS=2 THEN '租金未支付自动关闭'
                                  WHEN ff.STATUS=3 THEN '待发货（支付完成）'
                                  WHEN ff.STATUS=4 THEN '拣货中'
                                  WHEN ff.STATUS=5 THEN '已发货'
                                  WHEN ff.STATUS=6 THEN '交易成功'
                                  WHEN ff.STATUS=7 THEN '交易取消（客服取消）'
                                  WHEN ff.STATUS=8 THEN '已退货'
                                  WHEN ff.STATUS=9 THEN '结束'
                                  WHEN ff.STATUS=10 THEN '已收货'
                                  WHEN ff.STATUS=11 THEN '租金已付，保险未支付（流程限制，不会出现该状态）'
                                  WHEN ff.STATUS=12 THEN '保险已支付，租金未支付'
                                  WHEN ff.STATUS=13 THEN '待授权'
                                  WHEN ff.STATUS=14 THEN '维修中'
                                  WHEN ff.STATUS=15 THEN '退货中'
                                  WHEN ff.STATUS=22 THEN '保险未支付自动关闭'
                                  WHEN ff.STATUS=23 THEN '用户关闭订单'
                                  WHEN ff.STATUS=25 THEN '客服审核中'
                                  WHEN ff.STATUS=26 THEN '还机中'
                                  WHEN ff.STATUS=27 THEN '已买断'
                                  WHEN ff.STATUS=28 THEN '还机成功'
                                  WHEN ff.STATUS=29 THEN '还机失败'
                                  WHEN ff.STATUS=30 THEN '逾期'
                                  WHEN ff.STATUS=31 THEN '租金逾期'
                                  WHEN ff.STATUS=72 THEN '交易取消（客户取消）'
                                  ELSE NULL END) AS trade_status,
                            dd.approve_status
                            --,ee.address
                            from
                            (
                            select b.risk_user_id from dw_2.order_all a
                            left join dw_2.user_id_relate_risk b on a.user_id=b.user_id --拼接risk_user_id
                            where a.trade_no='20190505100230101434'
                            ) aa
                            left join 
                            (
                            select b.risk_user_id,a.* from dw_2.order_all a
                            left join dw_2.user_id_relate_risk b on a.user_id=b.user_id --拼接risk_user_id
                            ) bb
                            on aa.risk_user_id=bb.risk_user_id
                            left join dw_2.product_sku_main cc on bb.id_sku=cc.sku_id
                            left join dw_2.risk_after_sale_approve_result dd on bb.trade_no=dd.trade_no
                            left join stage.airentdbsetting_airent_tbl_trade_delivery_address ee on bb.trade_no=ee.trade_no
                            left join dw_2.order_life_status_cur ff on bb.trade_no=ff.trade_no
                            order by to_date(bb.dt_created) asc
                           '''
                              , conn_hive)


application_address_log = pd.read_sql('''
                            select 
                            bb.dt_created,
                            --bb.trade_no,bb.risk_user_id,bb.user_id,
                            (CASE WHEN ff.STATUS=0 THEN '初始状态'
                                  WHEN ff.STATUS=1 THEN '待支付（订单提交）'
                                  WHEN ff.STATUS=2 THEN '租金未支付自动关闭'
                                  WHEN ff.STATUS=3 THEN '待发货（支付完成）'
                                  WHEN ff.STATUS=4 THEN '拣货中'
                                  WHEN ff.STATUS=5 THEN '已发货'
                                  WHEN ff.STATUS=6 THEN '交易成功'
                                  WHEN ff.STATUS=7 THEN '交易取消（客服取消）'
                                  WHEN ff.STATUS=8 THEN '已退货'
                                  WHEN ff.STATUS=9 THEN '结束'
                                  WHEN ff.STATUS=10 THEN '已收货'
                                  WHEN ff.STATUS=11 THEN '租金已付，保险未支付（流程限制，不会出现该状态）'
                                  WHEN ff.STATUS=12 THEN '保险已支付，租金未支付'
                                  WHEN ff.STATUS=13 THEN '待授权'
                                  WHEN ff.STATUS=14 THEN '维修中'
                                  WHEN ff.STATUS=15 THEN '退货中'
                                  WHEN ff.STATUS=22 THEN '保险未支付自动关闭'
                                  WHEN ff.STATUS=23 THEN '用户关闭订单'
                                  WHEN ff.STATUS=25 THEN '客服审核中'
                                  WHEN ff.STATUS=26 THEN '还机中'
                                  WHEN ff.STATUS=27 THEN '已买断'
                                  WHEN ff.STATUS=28 THEN '还机成功'
                                  WHEN ff.STATUS=29 THEN '还机失败'
                                  WHEN ff.STATUS=30 THEN '逾期'
                                  WHEN ff.STATUS=31 THEN '租金逾期'
                                  WHEN ff.STATUS=72 THEN '交易取消（客户取消）'
                                  ELSE NULL END) AS trade_status,
                            dd.approve_status,
                            ee.address
                            from
                            (
                            select b.risk_user_id from dw_2.order_all a
                            left join dw_2.user_id_relate_risk b on a.user_id=b.user_id --拼接risk_user_id
                            where a.trade_no='20190505100230101434'
                            ) aa
                            left join 
                            (
                            select b.risk_user_id,a.* from dw_2.order_all a
                            left join dw_2.user_id_relate_risk b on a.user_id=b.user_id --拼接risk_user_id
                            ) bb
                            on aa.risk_user_id=bb.risk_user_id
                            left join dw_2.product_sku_main cc on bb.id_sku=cc.sku_id
                            left join dw_2.risk_after_sale_approve_result dd on bb.trade_no=dd.trade_no
                            left join stage.airentdbsetting_airent_tbl_trade_delivery_address ee on bb.trade_no=ee.trade_no
                            left join dw_2.order_life_status_cur ff on bb.trade_no=ff.trade_no
                            order by to_date(bb.dt_created) asc
                           '''
                                      , conn_hive)


evaluate_score_info = pd.read_sql('''
                            select 
                            bb.dt_evaluate,bb.score_card_id,bb.report_no,bb.evaluate_score,bb.evaluate_credit,bb.use_credit 
                            from 
                            (
                            select a.trade_no,a.user_id,b.risk_user_id from dw_2.order_all a
                            left join dw_2.user_id_relate_risk b on a.user_id=b.user_id 
                            where a.trade_no='20190505100230101434'
                            ) aa
                            left join dw_2.risk_report_log bb on aa.risk_user_id=bb.risk_user_id
                           '''
                                  , conn_hive)


zm_info = pd.read_sql('''
                    select 
                    bb.dt_created,bb.user_name,bb.zm_scope,bb.zm_grade,bb.zm_face,bb.zm_risk
                    from 
                    (
                    select a.trade_no,a.user_id,b.risk_user_id from dw_2.order_all a
                    left join dw_2.user_id_relate_risk b on a.user_id=b.user_id 
                    where a.trade_no='20190505100230101434'
                    ) aa
                    left join dw_2.user_detail_basic bb on aa.user_id=bb.user_id
                    '''
                      , conn_hive)


paipai_info = pd.read_sql('''
                    select 
                    bb.dt_created,bb.dt_expired,bb.score,bb.risk_rank,bb.probability,bb.consumer,bb.device,bb.social,bb.is_black,bb.is_alert
                    from 
                    (
                    select a.trade_no,a.user_id,b.risk_user_id from dw_2.order_all a
                    left join dw_2.user_id_relate_risk b on a.user_id=b.user_id 
                    where a.trade_no='20190505100230101434'
                    ) aa
                    left join stage.airentRiskDbSetting_xhj_risk_control_e_paipai_data bb on aa.risk_user_id=bb.user_id
                    '''
                          , conn_hive)


tongdun_info = pd.read_sql('''
                    select 
                    bb.dt_created,bb.dt_expired,bb.score
                    from 
                    (
                    select a.trade_no,a.user_id,b.risk_user_id from dw_2.order_all a
                    left join dw_2.user_id_relate_risk b on a.user_id=b.user_id 
                    where a.trade_no='20190505100230101434'
                    ) aa
                    left join stage.airentRiskDbSetting_xhj_risk_control_e_tongdun_data bb on aa.risk_user_id=bb.user_id
                    '''
                           , conn_hive)


suolun_info = pd.read_sql('''
                    select
                    bb.dt_created,
                    func.get_json_object(content,'$.user_name') as 名称,
                    func.get_json_object(content,'$.user_phone') as 手机号,
                    func.get_json_object(content,'$.user_idcard') as 身份证号,
                    func.get_json_object(content,'$.binding_idcards') as 绑定其他身份证信息,
                    func.get_json_object(content,'$.binding_phones') as 绑定其他手机号信息,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.age') as 年龄,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.gender') as 性别,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.birthday') as 生日,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.idcard_validate') as 是否有效身份证,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.idcard_province') as 身份证省份,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.idcard_city') as 身份证城市,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.idcard_region') as 身份证地区,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.record_idcard_days') as 身份证记录天数,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.last_appear_idcard') as 身份证最近出现时间,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.used_idcards_cnt') as 身份证关联数量,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.phone_operator') as 手机运营商,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.phone_province') as 手机归属地省份,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.phone_city') as 手机归属地城市,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.record_phone_days') as 手机号出现时间,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.last_appear_phone') as 手机号最近出现时间,
                    func.get_json_object(func.get_json_object(content,'$.user_basic'),'$.used_phones_cnt') as 手机号关联数量,
                    func.get_json_object(func.get_json_object(content,'$.risk_social_network'),'$.sn_score') as 葫芦分,
                    func.get_json_object(func.get_json_object(content,'$.risk_social_network'),'$.sn_order1_contacts_cnt') as 直接联系人,
                    func.get_json_object(func.get_json_object(content,'$.risk_social_network'),'$.sn_order1_blacklist_contacts_cnt') as 直接联系人在黑名单中数量,
                    func.get_json_object(func.get_json_object(content,'$.risk_social_network'),'$.sn_order2_blacklist_routers_cnt') as 间接联系人在黑名单中数量,
                    func.get_json_object(func.get_json_object(content,'$.risk_social_network'),'$.sn_order2_blacklist_contacts_cnt') as 认识间接黑人的直接联系人个数,
                    func.get_json_object(func.get_json_object(content,'$.risk_social_network'),'$.sn_order2_blacklist_routers_pct') as 认识间接黑人的直接联系人比例,
                    func.get_json_object(func.get_json_object(content,'$.risk_blacklist'),'$.idcard_in_blacklist') as 身份证是否命中黑名单,
                    func.get_json_object(func.get_json_object(content,'$.risk_blacklist'),'$.phone_in_blacklist') as 手机号是否命中黑名单,
                    func.get_json_object(func.get_json_object(content,'$.risk_blacklist'),'$.in_p2p_blacklist') as 是否命中P2P黑名单,
                    func.get_json_object(func.get_json_object(content,'$.risk_blacklist'),'$.in_court_blacklist') as 是否命中法院黑名单,
                    func.get_json_object(func.get_json_object(content,'$.risk_blacklist'),'$.in_bank_blacklist') as 是否命中银行黑名单,
                    func.get_json_object(func.get_json_object(content,'$.risk_blacklist'),'$.last_appear_idcard_in_blacklist') as 最近身份证出现在黑名单内时间,
                    func.get_json_object(func.get_json_object(content,'$.risk_blacklist'),'$.last_appear_phone_in_blacklist') as 最近手机号出现在黑名单内时间,
                    func.get_json_object(func.get_json_object(content,'$.history_org'),'$.online_cash_loan_cnt') as 线上现金贷出现次数,
                    func.get_json_object(func.get_json_object(content,'$.history_org'),'$.offline_cash_loan_cnt') as 线下现金贷出现次数,
                    func.get_json_object(func.get_json_object(content,'$.history_org'),'$.online_installment_cnt') as 线上分期消费出现次数,
                    func.get_json_object(func.get_json_object(content,'$.history_org'),'$.offline_installment_cnt') as 线下分期消费出现次数,
                    func.get_json_object(func.get_json_object(content,'$.history_org'),'$.payday_loan_cnt') as 小额贷款出现次数,
                    func.get_json_object(func.get_json_object(content,'$.history_org'),'$.credit_card_repayment_cnt') as 信用卡代还出现次数,
                    func.get_json_object(func.get_json_object(content,'$.history_org'),'$.others_cnt') as 其他次数,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.search_cnt') as 历史查询次数,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.search_cnt_recent_7_days') as 最近7填查询次数,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.search_cnt_recent_14_days') as 最近14填查询次数,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.search_cnt_recent_30_days') as 最近30填查询次数,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.search_cnt_recent_60_days') as 最近60填查询次数,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.search_cnt_recent_90_days') as 最近90填查询次数,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.search_cnt_recent_180_days') as 最近180填查询次数,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.org_cnt') as 历史机构查询数量,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.org_cnt_recent_7_days') as 最近7天机构查询数量,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.org_cnt_recent_14_days') as 最近14天机构查询数量,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.org_cnt_recent_30_days') as 最近30天机构查询数量,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.org_cnt_recent_60_days') as 最近60天机构查询数量,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.org_cnt_recent_90_days') as 最近90天机构查询数量,
                    func.get_json_object(func.get_json_object(content,'$.history_search'),'$.org_cnt_recent_180_days') as 最近180天机构查询数量
                    from
                    (
                    select a.trade_no,a.user_id,b.risk_user_id from dw_2.order_all a
                    left join dw_2.user_id_relate_risk b on a.user_id=b.user_id 
                    where a.trade_no='20190505100230101434'
                    ) aa
                    left join stage.airentRiskDbSetting_xhj_risk_control_e_blacklist_user_sauron bb on aa.risk_user_id=bb.uid
                    order by dt_created desc
                    '''
                          , conn_hive)


address_book = pd.read_sql('''
                    select
                    row_number() over(order by bb.dt_updated) as id, 
                    bb.name,bb.phone,bb.dt_updated
                    from 
                    (
                    select a.trade_no,a.user_id,b.risk_user_id from dw_2.order_all a
                    left join dw_2.user_id_relate_risk b on a.user_id=b.user_id 
                    where a.trade_no='20190505100230101434'
                    ) aa
                    left join stage.DataDbSetting_xhj_warehouse_user_address_book bb on aa.risk_user_id=bb.risk_id_user
                    '''
                           , conn_hive)


contact_info = pd.read_sql('''
                    select
                    bb.dt_created,bb.user_id,bb.contact_name,bb.contact_phone
                    from 
                    (
                    select a.trade_no,a.user_id,b.risk_user_id from dw_2.order_all a
                    left join dw_2.user_id_relate_risk b on a.user_id=b.user_id 
                    where a.trade_no='20190505100230101434'
                    ) aa
                    left join stage.airentriskdbsetting_xhj_risk_control_e_user_credit_contact bb on aa.risk_user_id=bb.user_id
                    order by dt_created desc
                    '''
                           , conn_hive)


log_info = pd.read_sql('''
                    select
                    bb.dt_created,bb.model_type,bb.android_id,bb.imei,bb.idfa,
                    bb.ip,bb.country,bb.province,bb.city,bb.isp,bb.type
                    from
                    (
                    select a.trade_no,a.user_id,b.risk_user_id from dw_2.order_all a
                    left join dw_2.user_id_relate_risk b on a.user_id=b.user_id 
                    where a.trade_no='20190505100230101434'
                    ) aa
                    left join dw_2.user_login_log bb on aa.user_id=bb.user_id
                    order by dt_created desc
                    '''
                       , conn_hive)


document = Document()
document.add_heading('调查报告', 0)

document.add_heading('基本信息', level=0)
for i in hive_base_info.columns:
    document.add_paragraph(str(i) + ' : ' + str(hive_base_info[i][0]), style='ListBullet')

document.add_heading('联系人信息', level=0)
for i in contact_info.columns:
    document.add_paragraph(str(i) + ' : ' + str(contact_info[i][0]), style='ListBullet')

document.add_heading('订单(还款|逾期)状态', level=0)
table = document.add_table(rows=len(sql_repay_info) + 1, cols=len(sql_repay_info.columns), style='Table Grid')
for i in range(0, len(sql_repay_info.columns)):  # 表标签
    cell = table.cell(0, i)
cell.text = str(sql_repay_info.columns[i])
for i in sql_repay_info.columns:
    for
x in range(0, len(sql_repay_info)):
# print(i,list(sql_repay_info.columns).index(i))
cell = table.cell(x + 1, list(sql_repay_info.columns).index(i))
# rint(sql_repay_info[i][x])
cell.text = str(sql_repay_info[i][x])

document.add_heading('催记', level=0)
table = document.add_table(rows=len(collect_note) + 1, cols=len(collect_note.columns), style='Table Grid')
for i in range(0, len(collect_note.columns)):  # 表标签
    cell = table.cell(0, i)
cell.text = str(collect_note.columns[i])
for i in collect_note.columns:
    for
x in range(0, len(collect_note)):
# print(i,list(sql_repay_info.columns).index(i))
cell = table.cell(x + 1, list(collect_note.columns).index(i))
# rint(sql_repay_info[i][x])
cell.text = str(collect_note[i][x])

document.add_heading('审核记录', level=0)
table = document.add_table(rows=len(approve_log) + 1, cols=len(approve_log.columns), style='Table Grid')
for i in range(0, len(approve_log.columns)):  # 表标签
    cell = table.cell(0, i)
cell.text = str(approve_log.columns[i])
for i in approve_log.columns:
    for
x in range(0, len(approve_log)):
# print(i,list(sql_repay_info.columns).index(i))
cell = table.cell(x + 1, list(approve_log.columns).index(i))
# rint(sql_repay_info[i][x])
cell.text = str(approve_log[i][x])

document.add_heading('历史申请记录', level=0)
table = document.add_table(rows=len(application_log) + 1, cols=len(application_log.columns), style='Table Grid')
for i in range(0, len(application_log.columns)):  # 表标签
    cell = table.cell(0, i)
cell.text = str(application_log.columns[i])
for i in application_log.columns:
    for
x in range(0, len(application_log)):
# print(i,list(sql_repay_info.columns).index(i))
cell = table.cell(x + 1, list(application_log.columns).index(i))
# rint(sql_repay_info[i][x])
cell.text = str(application_log[i][x])

document.add_heading('历史申请地址', level=0)
table = document.add_table(rows=len(application_address_log) + 1, cols=len(application_address_log.columns),
                           style='Table Grid')
for i in range(0, len(application_address_log.columns)):  # 表标签
    cell = table.cell(0, i)
cell.text = str(application_address_log.columns[i])
for i in application_address_log.columns:
    for
x in range(0, len(application_address_log)):
# print(i,list(sql_repay_info.columns).index(i))
cell = table.cell(x + 1, list(application_address_log.columns).index(i))
# rint(sql_repay_info[i][x])
cell.text = str(application_address_log[i][x])

document.add_heading('信用评估记录', level=0)
table = document.add_table(rows=len(evaluate_score_info) + 1, cols=len(evaluate_score_info.columns), style='Table Grid')
for i in range(0, len(evaluate_score_info.columns)):  # 表标签
    cell = table.cell(0, i)
cell.text = str(evaluate_score_info.columns[i])
for i in evaluate_score_info.columns:
    for
x in range(0, len(evaluate_score_info)):
# print(i,list(sql_repay_info.columns).index(i))
cell = table.cell(x + 1, list(evaluate_score_info.columns).index(i))
# rint(sql_repay_info[i][x])
cell.text = str(evaluate_score_info[i][x])

document.add_heading('芝麻信息', level=0)
for i in zm_info.columns:
    document.add_paragraph(str(i) + ' : ' + str(zm_info[i][0]), style='ListBullet')

document.add_heading('拍拍信息', level=0)
table = document.add_table(rows=len(paipai_info) + 1, cols=len(paipai_info.columns), style='Table Grid')
for i in range(0, len(paipai_info.columns)):  # 表标签
    cell = table.cell(0, i)
cell.text = str(paipai_info.columns[i])
for i in paipai_info.columns:
    for
x in range(0, len(paipai_info)):
# print(i,list(sql_repay_info.columns).index(i))
cell = table.cell(x + 1, list(paipai_info.columns).index(i))
# rint(sql_repay_info[i][x])
cell.text = str(paipai_info[i][x])

document.add_heading('同盾信息', level=0)
table = document.add_table(rows=len(tongdun_info) + 1, cols=len(tongdun_info.columns), style='Table Grid')
for i in range(0, len(tongdun_info.columns)):  # 表标签
    cell = table.cell(0, i)
cell.text = str(tongdun_info.columns[i])
for i in tongdun_info.columns:
    for
x in range(0, len(tongdun_info)):
# print(i,list(sql_repay_info.columns).index(i))
cell = table.cell(x + 1, list(tongdun_info.columns).index(i))
# rint(sql_repay_info[i][x])
cell.text = str(tongdun_info[i][x])

document.add_heading('葫芦索伦信息', level=0)
for i in suolun_info.columns:
    document.add_paragraph(str(i) + ' : ' + str(suolun_info[i][0]), style='ListBullet')

document.add_heading('通讯录', level=0)
table = document.add_table(rows=len(address_book) + 1, cols=len(address_book.columns), style='Table Grid')
for i in range(0, len(address_book.columns)):  # 表标签
    cell = table.cell(0, i)
cell.text = str(address_book.columns[i])
for i in address_book.columns:
    for
x in range(0, len(address_book)):
# print(i,list(sql_repay_info.columns).index(i))
cell = table.cell(x + 1, list(address_book.columns).index(i))
# rint(sql_repay_info[i][x])
cell.text = str(address_book[i][x])

document.add_heading('登陆记录', level=0)
table = document.add_table(rows=len(log_info) + 1, cols=len(log_info.columns), style='Table Grid')
for i in range(0, len(log_info.columns)):  # 表标签
    cell = table.cell(0, i)
cell.text = str(log_info.columns[i])
for i in log_info.columns:
    for
x in range(0, len(log_info)):
# print(i,list(sql_repay_info.columns).index(i))
cell = table.cell(x + 1, list(log_info.columns).index(i))
# rint(sql_repay_info[i][x])
cell.text = str(log_info[i][x])

document.save(r'C:\Users\jizeyuan\Desktop\risk_report.docx')

